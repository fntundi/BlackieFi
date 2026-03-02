"""
Knowledge AI Service - RAG-based document analysis using Emergent LLM
Supports documents, images, and videos for AI analysis
"""
import os
import uuid
from typing import Optional, List, Dict, Any
from pathlib import Path
from datetime import datetime, timezone
from dotenv import load_dotenv

load_dotenv()

# Configuration
UPLOAD_DIR = Path("/app/uploads/knowledge")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


class KnowledgeAIService:
    """Service for AI-powered document analysis using RAG"""
    
    SUPPORTED_TEXT_EXTENSIONS = ['txt', 'csv', 'md']
    SUPPORTED_DOC_EXTENSIONS = ['pdf', 'docx', 'xlsx']
    SUPPORTED_IMAGE_EXTENSIONS = ['png', 'jpg', 'jpeg', 'webp', 'gif', 'heic']
    SUPPORTED_VIDEO_EXTENSIONS = ['mp4', 'mov', 'webm', 'avi', 'mkv']
    
    MIME_TYPES = {
        'txt': 'text/plain',
        'csv': 'text/csv',
        'md': 'text/markdown',
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'webp': 'image/webp',
        'gif': 'image/gif',
        'heic': 'image/heic',
        'mp4': 'video/mp4',
        'mov': 'video/quicktime',
        'webm': 'video/webm',
        'avi': 'video/x-msvideo',
        'mkv': 'video/x-matroska',
    }
    
    def __init__(self):
        self.api_key = os.environ.get("EMERGENT_LLM_KEY")
        if not self.api_key:
            raise ValueError("EMERGENT_LLM_KEY not configured")
    
    def _get_file_extension(self, filename: str) -> str:
        """Get lowercase file extension"""
        return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    
    def _get_mime_type(self, extension: str) -> str:
        """Get MIME type from extension"""
        return self.MIME_TYPES.get(extension, 'application/octet-stream')
    
    async def _extract_text_content(self, file_path: Path, extension: str) -> str:
        """Extract text content from supported files including PDF, DOCX, XLSX"""
        try:
            # Plain text files
            if extension in self.SUPPORTED_TEXT_EXTENSIONS:
                with open(file_path, 'r', errors='ignore') as f:
                    content = f.read()
                return content[:50000]  # Limit to 50k chars
            
            # PDF files using PyMuPDF
            elif extension == 'pdf':
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(str(file_path))
                    text_parts = []
                    for page_num, page in enumerate(doc):
                        text = page.get_text()
                        if text.strip():
                            text_parts.append(f"[Page {page_num + 1}]\n{text}")
                    doc.close()
                    full_text = "\n\n".join(text_parts)
                    return full_text[:50000] if full_text else "No text content found in PDF"
                except Exception as e:
                    return f"PDF extraction error: {str(e)}"
            
            # DOCX files using python-docx
            elif extension == 'docx':
                try:
                    from docx import Document
                    doc = Document(str(file_path))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                paragraphs.append(row_text)
                    
                    full_text = "\n".join(paragraphs)
                    return full_text[:50000] if full_text else "No text content found in DOCX"
                except Exception as e:
                    return f"DOCX extraction error: {str(e)}"
            
            # XLSX files using openpyxl
            elif extension == 'xlsx':
                try:
                    from openpyxl import load_workbook
                    wb = load_workbook(str(file_path), data_only=True)
                    text_parts = []
                    
                    for sheet_name in wb.sheetnames:
                        sheet = wb[sheet_name]
                        sheet_data = [f"[Sheet: {sheet_name}]"]
                        
                        for row in sheet.iter_rows(values_only=True):
                            row_values = [str(cell) if cell is not None else '' for cell in row]
                            if any(v.strip() for v in row_values):
                                sheet_data.append(' | '.join(row_values))
                        
                        if len(sheet_data) > 1:  # Has data beyond header
                            text_parts.append("\n".join(sheet_data))
                    
                    wb.close()
                    full_text = "\n\n".join(text_parts)
                    return full_text[:50000] if full_text else "No data found in XLSX"
                except Exception as e:
                    return f"XLSX extraction error: {str(e)}"
            
            return ""
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    
    async def analyze_document_with_gemini(
        self,
        file_path: Path,
        query: str,
        session_id: str,
        file_type: str,
        system_context: str = ""
    ) -> Dict[str, Any]:
        """
        Analyze a document using Gemini's multimodal capabilities.
        Gemini supports file attachments for PDFs, images, text files, etc.
        """
        try:
            from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
            
            extension = self._get_file_extension(str(file_path))
            mime_type = self._get_mime_type(extension)
            
            # Build system message for financial analysis
            system_message = f"""You are an expert financial analyst AI assistant for BlackieFi, 
an institutional-grade wealth management platform. Your role is to analyze uploaded documents, 
images, and financial data to provide actionable insights.

{system_context}

When analyzing financial documents:
1. Identify key financial metrics and figures
2. Highlight important trends or patterns
3. Note any risks or concerns
4. Provide actionable recommendations
5. Be specific with numbers and data points

Format your response clearly with sections for:
- Key Findings
- Important Metrics (if applicable)
- Analysis & Insights
- Recommendations"""

            # Initialize chat with Gemini (required for file attachments)
            chat = LlmChat(
                api_key=self.api_key,
                session_id=session_id,
                system_message=system_message
            ).with_model("gemini", "gemini-2.5-flash")
            
            # Create file attachment
            file_content = FileContentWithMimeType(
                file_path=str(file_path),
                mime_type=mime_type
            )
            
            # Build the query message
            analysis_query = query or f"Please analyze this {file_type} and provide key financial insights."
            
            user_message = UserMessage(
                text=analysis_query,
                file_contents=[file_content]
            )
            
            # Send message and get response
            response = await chat.send_message(user_message)
            
            return {
                "success": True,
                "analysis": response,
                "model_used": "gemini-2.5-flash",
                "file_type": file_type,
                "query": analysis_query
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "analysis": f"Analysis failed: {str(e)}"
            }
    
    async def analyze_text_document(
        self,
        file_path: Path,
        query: str,
        session_id: str,
        system_context: str = ""
    ) -> Dict[str, Any]:
        """
        Analyze a text document using GPT-5.2 for text-based analysis.
        Falls back to this for simple text files where Gemini attachment isn't needed.
        """
        try:
            from emergentintegrations.llm.chat import LlmChat, UserMessage
            
            extension = self._get_file_extension(str(file_path))
            
            # Extract text content
            content = await self._extract_text_content(file_path, extension)
            if not content:
                return {
                    "success": False,
                    "error": "Could not extract text from document"
                }
            
            system_message = f"""You are an expert financial analyst AI assistant for BlackieFi.
Analyze the provided document content and provide actionable financial insights.

{system_context}

Focus on:
1. Key financial data and metrics
2. Trends and patterns
3. Risks and opportunities
4. Actionable recommendations"""

            chat = LlmChat(
                api_key=self.api_key,
                session_id=session_id,
                system_message=system_message
            ).with_model("openai", "gpt-5.2")
            
            analysis_query = query or "Please analyze this document and provide key financial insights."
            full_prompt = f"Document Content:\n```\n{content}\n```\n\nAnalysis Request: {analysis_query}"
            
            user_message = UserMessage(text=full_prompt)
            response = await chat.send_message(user_message)
            
            return {
                "success": True,
                "analysis": response,
                "model_used": "gpt-5.2",
                "file_type": "document",
                "query": analysis_query
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "analysis": f"Analysis failed: {str(e)}"
            }
    
    async def analyze_document(
        self,
        doc_metadata: Dict[str, Any],
        query: Optional[str] = None,
        user_id: str = "default"
    ) -> Dict[str, Any]:
        """
        Main entry point for document analysis.
        Routes to appropriate analysis method based on file type.
        """
        file_path = UPLOAD_DIR / doc_metadata["filename"]
        if not file_path.exists():
            return {
                "success": False,
                "error": "File not found on disk"
            }
        
        extension = self._get_file_extension(doc_metadata["filename"])
        file_type = doc_metadata.get("file_type", "document")
        session_id = f"knowledge-{user_id}-{doc_metadata.get('id', 'default')}"
        
        # Build context from document metadata
        system_context = f"""Document Information:
- Filename: {doc_metadata.get('original_filename', 'Unknown')}
- Type: {file_type}
- Size: {doc_metadata.get('file_size', 0)} bytes
- Description: {doc_metadata.get('description', 'No description provided')}
- Tags: {', '.join(doc_metadata.get('tags', []))}"""
        
        # Route to appropriate analysis method
        if file_type in ['image', 'video']:
            # Use Gemini for images and videos (visual analysis required)
            return await self.analyze_document_with_gemini(
                file_path=file_path,
                query=query,
                session_id=session_id,
                file_type=file_type,
                system_context=system_context
            )
        elif extension in ['pdf', 'docx', 'xlsx'] + self.SUPPORTED_TEXT_EXTENSIONS:
            # Use text extraction + GPT for documents with extractable text
            # This provides more reliable parsing than Gemini for complex documents
            return await self.analyze_text_document(
                file_path=file_path,
                query=query,
                session_id=session_id,
                system_context=system_context
            )
        else:
            # Try Gemini as fallback for other formats
            return await self.analyze_document_with_gemini(
                file_path=file_path,
                query=query,
                session_id=session_id,
                file_type=file_type,
                system_context=system_context
            )
    
    async def chat_with_knowledge_base(
        self,
        documents: List[Dict[str, Any]],
        query: str,
        user_id: str,
        chat_history: List[Dict[str, str]] = None
    ) -> Dict[str, Any]:
        """
        Chat with the knowledge base - answer questions based on uploaded documents.
        This creates a conversational RAG experience.
        """
        try:
            from emergentintegrations.llm.chat import LlmChat, UserMessage
            
            # Build context from documents
            doc_summaries = []
            for doc in documents[:10]:  # Limit to 10 most relevant docs
                doc_summaries.append(
                    f"- {doc.get('original_filename', 'Unknown')}: {doc.get('description', 'No description')} "
                    f"(Type: {doc.get('file_type', 'document')}, Tags: {', '.join(doc.get('tags', []))})"
                )
            
            doc_context = "\n".join(doc_summaries) if doc_summaries else "No documents available"
            
            system_message = f"""You are an AI financial research assistant for BlackieFi. 
You have access to the user's uploaded knowledge base of financial documents.

Available Documents in Knowledge Base:
{doc_context}

Your role:
1. Answer questions about the user's uploaded documents
2. Provide financial insights and analysis
3. Help with research and due diligence
4. Suggest relevant documents when applicable

If asked about specific documents, refer to them by name and provide relevant insights.
If you cannot answer based on the available documents, say so clearly."""

            session_id = f"knowledge-chat-{user_id}"
            
            chat = LlmChat(
                api_key=self.api_key,
                session_id=session_id,
                system_message=system_message
            ).with_model("openai", "gpt-5.2")
            
            # Build conversation context
            context_prompt = ""
            if chat_history:
                for msg in chat_history[-5:]:  # Last 5 messages for context
                    role = "User" if msg["role"] == "user" else "Assistant"
                    context_prompt += f"{role}: {msg['content']}\n"
            
            full_query = f"{context_prompt}\nUser: {query}" if context_prompt else query
            
            user_message = UserMessage(text=full_query)
            response = await chat.send_message(user_message)
            
            return {
                "success": True,
                "response": response,
                "model_used": "gpt-5.2",
                "documents_referenced": len(documents)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "response": f"Chat failed: {str(e)}"
            }


# Singleton instance
_knowledge_ai_service: Optional[KnowledgeAIService] = None

def get_knowledge_ai_service() -> KnowledgeAIService:
    """Get or create Knowledge AI service instance"""
    global _knowledge_ai_service
    if _knowledge_ai_service is None:
        _knowledge_ai_service = KnowledgeAIService()
    return _knowledge_ai_service
